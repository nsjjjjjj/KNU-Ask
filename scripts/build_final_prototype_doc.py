from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# compact_reference_guide preset.
# Named overrides: Noto Sans KR for Korean glyph coverage; memo_masthead opening;
# 27 pt title; compact table body text (8.5-9 pt) for genuine matrix data.
FONT = "Noto Sans KR"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "111827"
MUTED = "667085"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "147D64"
PALE_GREEN = "EAF7F2"
GOLD = "7A5A00"
PALE_GOLD = "FFF8E8"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), FONT)
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, bold=False, color=INK):
    style.font.name = FONT
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), FONT)
    r_fonts.set(qn("w:hint"), "eastAsia")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"table widths must total {CONTENT_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def add_numbering(doc, *, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT)
        r_fonts.set(qn("w:hAnsi"), FONT)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_number(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, cached, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_paragraph(paragraph, fill, border_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullet(doc, text, bullet_id):
    p = doc.add_paragraph()
    apply_number(p, bullet_id)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_numbered(doc, text, number_id):
    p = doc.add_paragraph()
    apply_number(p, number_id)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{label}  ")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(value)
    set_run_font(r, color=INK)
    return p


def add_callout(doc, label, text, *, color=BLUE, fill=CALLOUT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, fill, color)
    r = p.add_run(f"{label}  ")
    set_run_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_table(doc, headers, rows, widths, *, font_size=9.0, first_col_center=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header = table.rows[0]
    mark_repeat_header(header)
    for cell, value in zip(header.cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(value)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for values in rows:
        row = table.add_row()
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_col_center and index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_picture(doc, path, width, title, descr, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("title", title)
    picture._inline.docPr.set("descr", descr)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(0)
    r = cap.add_run(caption)
    set_run_font(r, size=8.5, color=MUTED, italic=True)


def build(output: Path, screenshot: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, False, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, size, True, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("강냉이 에스크  |  프로토타입 완성본")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(footer, "PAGE")

    bullet_id = add_numbering(doc, bullet=True)
    number_id = add_numbering(doc, bullet=False)

    # Page 1: product overview
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    r = kicker.add_run("PROTOTYPE FINAL BRIEF")
    set_run_font(r, size=9.5, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("강냉이 에스크")
    set_run_font(r, size=27, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("AI 학생지원 챗봇 프로토타입 완성 문서")
    set_run_font(r, size=13.5, color=MUTED, italic=True)

    add_label_value(doc, "작성", "김소연 · 박재희 · 노성준")
    add_label_value(doc, "핵심 사용자", "강남대학교 학생")
    add_label_value(doc, "기준일", "2026년 7월 21일")
    consensus = doc.add_paragraph()
    consensus.paragraph_format.space_after = Pt(9)
    r = consensus.add_run("회의 반영 기준  ")
    set_run_font(r, bold=True, color=NAVY)
    r = consensus.add_run("각 주제에서 세 구성원의 의견 뒤 작성된 ‘최종 정리’와 마지막 취합 메시지를 우선 반영")
    set_run_font(r, color=INK)

    add_callout(
        doc,
        "한 문장 설명",
        "학교 공식 공지와 FAQ를 바탕으로 학생에게 필요한 신청 정보·담당 부서·다음 행동을 한 화면에 정리해 주는 공개정보 안내 챗봇입니다.",
        color=GREEN,
        fill=PALE_GREEN,
    )

    add_heading(doc, "1. 서비스 개요", 1)
    add_heading(doc, "1.1 해결하려는 문제", 2)
    add_body(
        doc,
        "학생은 장학금·등록금·수강·휴복학 등 정보를 찾기 위해 여러 공지와 안내 페이지를 확인해야 합니다. 담당자는 이미 공개된 내용을 전화나 방문 문의로 반복 설명하고, 잘못 찾아온 문의를 다시 전달하는 부담을 겪고 있습니다.",
    )

    add_heading(doc, "1.2 목표와 기대 효과", 2)
    add_bullet(doc, "학생: 신청 대상·기간·준비 서류·처리 순서와 원문을 한 화면에서 확인", bullet_id)
    add_bullet(doc, "담당자: 단순 반복 문의를 줄이고 실제 판단이 필요한 문의에 집중", bullet_id)
    add_bullet(doc, "서비스: 근거가 부족할 때 추측하지 않고 추가 질문·원문·담당 부서 확인으로 전환", bullet_id)

    # Page 2 begins with the visual, so the overview page stays coherent and
    # the screenshot does not become a mostly empty overflow page.
    page_break(doc)
    if screenshot.exists():
        add_picture(
            doc,
            screenshot,
            5.45,
            "강냉이 에스크 답변 화면",
            "수강신청 일정 질문에 공식 공지 기반 요약, 일정, 담당 부서와 원문 확인 기능이 표시된 화면",
            "그림 1. 질문에 대한 공지 기반 답변 화면",
        )

    # Requirements and states continue directly below the visual.
    add_heading(doc, "2. 최종 사용자 요구사항", 1)
    add_heading(doc, "2.1 입력·결과·다음 행동", 2)
    add_table(
        doc,
        ["구분", "최종 합의 내용"],
        [
            ("입력값", "학생이 작성한 질문, 문의 분야, 자주 묻는 질문(FAQ)"),
            ("필수값", "질문 입력 또는 항목 선택 중 한 가지"),
            ("핵심 결과", "AI 답변과 공지 요약, 신청 대상·기간·준비 서류·처리 순서, 담당 부서·연락처, 원문 공지"),
            ("다음 행동", "추가 질문을 입력하거나 원문을 확인하고, 해결되지 않으면 담당 부서에 문의"),
            ("종료", "채팅 종료 버튼 선택 후 종료 여부를 확인하고 대화 화면을 초기화"),
        ],
        [1750, 7610],
        font_size=9.2,
        first_col_center=True,
    )

    add_heading(doc, "2.2 기본 이용 흐름", 2)
    flow_number_id = add_numbering(doc, bullet=False)
    for text in (
        "학교 홈페이지 또는 공유된 데모 주소에서 ‘강냉이 에스크’를 실행합니다.",
        "질문을 직접 입력하거나 문의 분야·FAQ를 선택합니다.",
        "서비스가 관련 공지와 FAQ를 검색해 질문에 맞는 답변을 구성합니다.",
        "학생은 요약, 대상, 기간, 준비 서류, 처리 순서, 담당 부서·연락처와 원문을 확인합니다.",
        "정보가 부족하면 추가 질문을 입력하고, 필요하면 원문 확인 또는 담당 부서 문의로 이동합니다.",
    ):
        add_numbered(doc, text, flow_number_id)

    add_heading(doc, "2.3 화면 상태별 동작", 2)
    add_table(
        doc,
        ["상태", "표시 내용", "사용자 처리"],
        [
            ("입력 전", "질문 입력창, 문의 분야, FAQ, 예시 질문", "질문 또는 항목 선택"),
            ("답변", "AI 답변, 공지 요약, 대상·기간·서류·절차, 담당 부서·연락처, 원문", "추가 질문·원문 확인·담당 부서 문의"),
            ("빈 입력", "현재 UI는 전송 버튼을 비활성화해 빈 질문 전송을 차단", "질문 입력 또는 항목 선택"),
            ("근거 없음", "‘공식 공지와 검수 FAQ에서 근거를 찾지 못했습니다’라고 안내", "학년도·학기·대상 보완 또는 담당 부서 확인"),
            ("종료", "‘문의를 종료하시겠습니까?’ 확인 모달", "계속 문의 또는 종료"),
        ],
        [1350, 5200, 2810],
        font_size=8.4,
        first_col_center=True,
    )

    # Scope follows naturally after the state matrix to avoid a sparse
    # continuation page when the table spans a page boundary.
    add_heading(doc, "3. 기능 범위와 답변 원칙", 1)
    add_heading(doc, "3.1 회의에서 합의한 우선순위", 2)
    add_table(
        doc,
        ["구분", "합의 기능", "최종 구현·범위"],
        [
            ("P0", "질문 입력·문의 분야 선택, 공지·FAQ 기반 답변, 담당 부서 안내, 추가 질문", "첫 버전 핵심 흐름으로 구현"),
            ("P1", "관련 공지 원문 링크, 공지 세부사항과 처리 순서", "원문 연결 구현, 세부 구조화는 대표 범위에서 제공"),
            ("P2", "새 공지 자동 반영, 맞춤 추천·알림·문의 통계, 품질 검증 고도화", "사용자 기능 범위에서는 제외; 수집 파이프라인은 운영 확장용으로 준비"),
        ],
        [1050, 4010, 4300],
        font_size=8.7,
        first_col_center=True,
    )

    add_callout(
        doc,
        "첫 버전 최종 범위",
        "포함: 채팅 질문, 문의 분야, FAQ, AI 답변, 담당 부서·연락처, 원문 공지 확인. 제외: 담당 부서로 즉시 전화 연결, 세부 카테고리, 자동 채팅 종료, 문의 원문 저장, 개인 맞춤 추천과 알림.",
        color=BLUE,
        fill=PALE_BLUE,
    )

    add_heading(doc, "3.2 답변에 포함하는 정보", 2)
    for text in (
        "학생이 지금 알아야 할 핵심 요약",
        "신청 대상과 제외 조건, 학년도·학기 등 적용 조건",
        "신청·제출·납부 기간과 예비·변경 등 중요 일정",
        "준비 서류, 비용, 장소·온라인 경로와 순서가 있는 처리 절차",
        "담당 부서, 담당 업무, 담당자·전화번호·운영시간(확인 가능한 범위)",
        "관련 공식 공지 원문, 신청 링크, 확인 기준 시각과 다음 행동",
    ):
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "3.3 공식 근거와 안전장치", 2)
    for text in (
        "근거 우선순위는 공지 원문·첨부파일 → 학사안내·공식 FAQ → 공식 직원 연락처입니다.",
        "공지에 연락처가 없을 때만 공식 직원 연락처에서 ‘부서 + 담당 업무’를 함께 대조해 보완합니다.",
        "원문에 없는 기간·절차·링크·연락처는 만들지 않고 ‘확인 필요’ 또는 담당 부서 문의로 전환합니다.",
        "개인정보와 대화 원문은 장기 저장하지 않으며 학번·전화번호·이메일 입력을 화면에서 차단합니다.",
    ):
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "3.4 데이터 처리 흐름", 2)
    add_body(
        doc,
        "학교 공지·학사안내·FAQ·직원 연락처의 원문과 첨부를 수집하고, 기간·대상·절차·연락처를 구조화합니다. 질문 시 키워드와 의미 검색으로 관련 근거를 찾고, 확인된 값만 답변 카드에 표시하며 공식 원문을 함께 제공합니다. PDF·이미지·Office 문서는 텍스트 추출 또는 OCR 결과와 원문 위치를 보존합니다.",
    )

    # Implementation continues without a forced break so available page space
    # is used before Word starts the next page.
    add_heading(doc, "4. 구현 구성과 완료 상태", 1)
    add_heading(doc, "4.1 최종 구현 구성", 2)
    add_table(
        doc,
        ["영역", "구성", "현재 역할"],
        [
            ("화면", "React 18 · TypeScript · Tailwind CSS · Vite", "채팅, 분야·FAQ, 답변 카드, 종료 확인"),
            ("서버", "FastAPI · Python 3.12 · SQLAlchemy", "질문 분석, 검색, 답변, 크롤러·관리 API"),
            ("데이터", "PostgreSQL 16 · pgvector", "공지 메타데이터, 구조화 정보, 키워드·의미 검색"),
            ("AI", "조건부 OpenAI/Gemini/Ollama + 결정적 응답", "복합 질문 답변, 공지 구조화, 임베딩"),
            ("수집", "증분 크롤링 · 변경 감지 · OCR/문서 추출", "새 공지와 변경 공지만 재처리"),
            ("검증", "pytest 67개 · 프런트엔드 7개 테스트 구성", "API, 검색, 안전 실패, 개인정보 차단, 주요 UI 흐름"),
        ],
        [1250, 3420, 4690],
        font_size=8.3,
        first_col_center=True,
    )

    add_callout(
        doc,
        "구현안 확정",
        "회의 중에는 Firebase 가상 데이터 방식이 논의됐지만, 최종 PRD와 현재 저장소는 React + FastAPI + PostgreSQL + pgvector 구조를 사용합니다. 문의 분야 선택은 미리 작성한 답변 대신 분야별 수집 공지를 보여주는 방식으로 구체화했습니다.",
        color=GOLD,
        fill=PALE_GOLD,
    )

    add_heading(doc, "4.2 완료 기준과 현재 상태", 2)
    add_table(
        doc,
        ["완료 기준", "현재 상태"],
        [
            ("질문을 입력하면 AI 답변을 받을 수 있음", "구현"),
            ("문의 분야·FAQ 선택으로 시작할 수 있음", "구현"),
            ("공지 요약과 담당 부서·연락처가 함께 표시됨", "구현"),
            ("관련 공지 원문을 확인할 수 있음", "구현"),
            ("근거가 없으면 추측하지 않고 질문 보완·담당 부서 확인을 안내함", "구현"),
            ("빈 입력 처리", "전송 버튼 비활성화 방식으로 구현"),
            ("채팅 종료 전 확인 및 대화 초기화", "구현"),
        ],
        [6760, 2600],
        font_size=8.6,
        first_col_center=False,
    )

    add_heading(doc, "5. 데모 실행", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("접속 주소  ")
    set_run_font(r, bold=True, color=NAVY)
    add_hyperlink(p, "프로토타입 접속하기", "https://lung-applicants-connectors-benefits.trycloudflare.com")
    r = p.add_run("  (https://lung-applicants-connectors-benefits.trycloudflare.com)")
    set_run_font(r, size=9.5, color=MUTED)
    demo_number_id = add_numbering(doc, bullet=False)
    for text in (
        "첫 화면에서 질문 입력, 문의 분야, FAQ의 세 가지 진입 방식을 확인합니다.",
        "‘2026학년도 2학기 수강신청 일정 알려줘’를 입력해 일정·담당 부서·원문을 확인합니다.",
        "‘휴학 신청 방법 알려줘’를 입력해 준비 서류와 순서형 처리 절차를 확인합니다.",
        "근거가 없는 질문으로 추측하지 않는 안내를 확인합니다.",
        "채팅 종료 버튼으로 종료 확인과 초기화 흐름을 확인합니다.",
    ):
        add_numbered(doc, text, demo_number_id)

    add_heading(doc, "6. 현재 제한사항과 운영 전 확인", 1)
    for text in (
        "이 문서는 공식 대학 서비스가 아닌 프로토타입을 설명합니다.",
        "샘플 모드의 일부 전화번호와 공지 URL은 시연용일 수 있으므로 실제 행동 전 공식 원문을 확인해야 합니다.",
        "학생 인증·개인 학사정보·담당 부서 즉시 전화 연결·상담원 이관은 포함하지 않습니다.",
        "TryCloudflare 주소는 임시 데모 주소이므로 장기 배포 전 고정 도메인과 부서별 데이터 검수 책임을 확정해야 합니다.",
    ):
        add_bullet(doc, text, bullet_id)

    core = doc.core_properties
    core.title = "강냉이 에스크 AI 학생지원 챗봇 프로토타입 완성본"
    core.subject = "PRD 및 팀 회의 최종 취합 의견 반영"
    core.author = "김소연, 박재희, 노성준"
    core.comments = "각 주제의 최종 정리와 마지막 취합 메시지를 우선 반영"
    core.keywords = "강냉이 에스크, 강남대학교, AI 챗봇, 프로토타입, PRD"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--screenshot", type=Path, required=True)
    args = parser.parse_args()
    build(args.output, args.screenshot)


if __name__ == "__main__":
    main()
