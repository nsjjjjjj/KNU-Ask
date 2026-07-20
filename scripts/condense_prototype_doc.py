from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt, RGBColor


FONT = "Arial Unicode MS"
BLUE = "2E74B5"
DARK = "243447"
MUTED = "667085"
LIGHT_BLUE = "EEF5FB"
GREEN = "178A72"
LIGHT_GREEN = "EAF7F3"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, bold=False, color=DARK):
    style.font.name = FONT
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        p_pr.append(borders)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_step(doc, number, title, detail):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{number}. {title}  ")
    set_run_font(r, size=10.5, bold=True, color=BLUE)
    r = p.add_run(detail)
    set_run_font(r, size=10.5, color=DARK)
    return p


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([fonts, color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([run_properties, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_access_step(doc, number, title, before_url, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{number}. {title}  ")
    set_run_font(r, size=10.5, bold=True, color=BLUE)
    r = p.add_run(before_url)
    set_run_font(r, size=10.5, color=DARK)
    add_hyperlink(p, url, url)
    return p


def build(output_path: Path, screenshot_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    title_style = doc.styles["Title"]
    set_style_font(title_style, 28, True, DARK)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(3)

    subtitle_style = doc.styles["Subtitle"]
    set_style_font(subtitle_style, 13, False, MUTED)
    subtitle_style.paragraph_format.space_after = Pt(10)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, True, BLUE)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 12, True, DARK)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    list_style = doc.styles["List Bullet"]
    set_style_font(list_style, 10.5, False, DARK)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("강냉이 에스크 · 프로토타입")
    set_run_font(hr, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(fp, "PAGE")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("강냉이 에스크")
    set_run_font(r, size=28, bold=True, color=DARK)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("AI 학생지원 챗봇 프로토타입 소개")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    r = meta.add_run("대상  ")
    set_run_font(r, size=9.5, bold=True, color=BLUE)
    r = meta.add_run("강남대학교 학생·학내 실무자    ")
    set_run_font(r, size=9.5, color=MUTED)
    r = meta.add_run("기준일  ")
    set_run_font(r, size=9.5, bold=True, color=BLUE)
    r = meta.add_run("2026. 7. 21.")
    set_run_font(r, size=9.5, color=MUTED)

    lead = doc.add_paragraph()
    lead.paragraph_format.left_indent = Cm(0.18)
    lead.paragraph_format.right_indent = Cm(0.1)
    lead.paragraph_format.space_before = Pt(2)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.22
    shade_paragraph(lead, LIGHT_GREEN, GREEN)
    r = lead.add_run(
        "학생이 필요한 담당 부서와 지원 정보를 더 빠르게 찾도록 돕고, 반복 문의에 대한 담당자의 응대 부담을 줄입니다."
    )
    set_run_font(r, size=11, bold=True, color=DARK)

    doc.add_heading("주요 기능", level=1)
    for item in (
        "채팅 질문, 문의 분야, 자주 묻는 질문 중 원하는 방식으로 시작",
        "관련 공지와 FAQ를 바탕으로 질문에 맞는 답변 제공",
        "신청 대상·기간·준비 서류·처리 순서 등 핵심 내용 정리",
        "담당 부서와 연락처, 관련 공지 원문 링크 안내",
        "답변이 부족하면 추가 질문을 입력하고, 해결되지 않으면 담당 부서로 문의",
    ):
        add_bullet(doc, item)

    picture_p = doc.add_paragraph()
    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_p.paragraph_format.space_before = Pt(8)
    picture_p.paragraph_format.space_after = Pt(2)
    picture = picture_p.add_run().add_picture(str(screenshot_path), width=Inches(6.2))
    picture._inline.docPr.set("title", "강냉이 에스크 프로토타입 화면")
    picture._inline.docPr.set(
        "descr", "수강신청 일정 질문에 공식 공지 기반 답변과 일정 카드가 표시된 챗봇 화면"
    )
    caption = doc.add_paragraph("프로토타입 답변 화면 예시")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(0)
    set_run_font(caption.runs[0], size=8.5, color=MUTED)

    doc.add_page_break()

    doc.add_heading("이용 흐름", level=1)
    demo_url = "https://lung-applicants-connectors-benefits.trycloudflare.com"
    add_access_step(doc, 1, "접속", "프로토타입 주소: ", demo_url)
    add_step(doc, 2, "질문 선택", "직접 질문하거나 문의 분야·자주 묻는 질문 선택")
    add_step(doc, 3, "답변 확인", "AI 답변, 공지 요약, 담당 부서 연락처, 원문 공지 확인")
    add_step(doc, 4, "다음 행동", "추가 질문을 입력하거나 원문을 확인하고, 필요하면 담당 부서에 문의")
    add_step(doc, 5, "종료", "채팅 종료 버튼을 누른 뒤 종료 여부 확인")

    doc.add_heading("화면별 동작", level=1)
    for item in (
        "입력 전: 질문 입력창, 문의 분야 선택, 자주 묻는 질문을 표시합니다.",
        "답변: 신청 대상·기간·준비 서류·처리 순서와 담당 부서·연락처·관련 공지를 보여줍니다.",
        "빈 입력: ‘질문을 입력하거나 문의 분야를 선택해 주세요’라고 안내합니다.",
        "정보 없음 또는 오류: 질문을 다시 입력하거나 담당 부서에 문의하도록 안내합니다.",
    ):
        add_bullet(doc, item)

    doc.add_heading("첫 버전 범위", level=1)
    included = doc.add_paragraph()
    included.paragraph_format.space_after = Pt(5)
    r = included.add_run("포함  ")
    set_run_font(r, size=10.5, bold=True, color=GREEN)
    r = included.add_run(
        "채팅 입력, 문의 분야 선택, FAQ, AI 답변, 담당 부서·연락처 안내, 원문 공지 확인"
    )
    set_run_font(r, size=10.5, color=DARK)

    excluded = doc.add_paragraph()
    excluded.paragraph_format.space_after = Pt(7)
    r = excluded.add_run("제외  ")
    set_run_font(r, size=10.5, bold=True, color=MUTED)
    r = excluded.add_run(
        "담당 부서로 바로 전화 연결, 새 공지·FAQ 자동 수집 및 반영, 세부 카테고리, 자동 채팅 종료, 문의 내용 저장"
    )
    set_run_font(r, size=10.5, color=DARK)

    scope_note = doc.add_paragraph()
    scope_note.paragraph_format.space_before = Pt(4)
    scope_note.paragraph_format.space_after = Pt(0)
    shade_paragraph(scope_note, LIGHT_BLUE, BLUE)
    r = scope_note.add_run("프로토타입 완료 기준  ")
    set_run_font(r, size=10, bold=True, color=BLUE)
    r = scope_note.add_run(
        "질문에 대한 답변과 공지 요약·담당 부서 안내가 함께 표시되고, 관련 공지 원문을 확인할 수 있는 상태입니다."
    )
    set_run_font(r, size=10, color=DARK)

    core = doc.core_properties
    core.title = "강냉이 에스크 AI 학생지원 챗봇 프로토타입 소개"
    core.subject = "회의 최종 정리 반영본"
    core.author = "강냉이 에스크 프로젝트팀"
    core.comments = "프로토타입 회의의 최종 합의 내용과 화면 흐름을 중심으로 정리"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--screenshot", required=True, type=Path)
    args = parser.parse_args()
    build(args.output, args.screenshot)


if __name__ == "__main__":
    main()
