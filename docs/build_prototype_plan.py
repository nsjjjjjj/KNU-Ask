from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "강냉이_에스크_실무자_피드백용_프로토타입_기획서.docx"
SCREENSHOT = ROOT / "assets" / "demo-answer.png"

# standard_business_brief preset + memo_masthead header pattern.
# Named typography override: Noto Sans KR preserves Korean glyphs in the
# macOS LibreOffice render path used for submission QA.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
BLACK = "111827"
GREEN = "147D64"
GOLD = "7A5A00"
RED = "9B1C1C"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Noto Sans KR"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), "Noto Sans KR")
    r_fonts.set(qn("w:hAnsi"), "Noto Sans KR")
    r_fonts.set(qn("w:eastAsia"), "Noto Sans KR")
    r_fonts.set(qn("w:cs"), "Noto Sans KR")
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def add_custom_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    nums = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(nums, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_number(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_hyperlink(paragraph, text, url, color=BLUE):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), color)
    r_pr.append(clr)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("강냉이 에스크  |  ")
    set_run_font(prefix, size=9, color=MUTED)
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=BLACK)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, color=BLACK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=BLACK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, color=BLACK)
    return p


def add_numbered(doc, text, num_id):
    p = doc.add_paragraph()
    apply_number(p, num_id)
    run = p.add_run(text)
    set_run_font(run, color=BLACK)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, bold=True, color=NAVY)
    value_run = p.add_run(value)
    set_run_font(value_run, color=BLACK)
    return p


def add_callout(doc, label, text, color=BLUE, fill=CALLOUT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.1)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    lead = p.add_run(f"{label}  ")
    set_run_font(lead, bold=True, color=color)
    body = p.add_run(text)
    set_run_font(body, color=BLACK)
    return p


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header = table.rows[0]
    mark_repeat_header(header)
    for cell, value in zip(header.cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for index, (cell, value) in enumerate(zip(row.cells, row_values)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Noto Sans KR"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
normal._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans KR")
normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

heading_tokens = {
    1: (16, BLUE, 16, 8),
    2: (13, BLUE, 12, 6),
    3: (12, DARK_BLUE, 8, 4),
}
for level, (size, color, before, after) in heading_tokens.items():
    style = styles[f"Heading {level}"]
    style.font.name = "Noto Sans KR"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
    style._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans KR")
    style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

list_style = styles["List Bullet"]
list_style.font.name = "Noto Sans KR"
list_style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
list_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
list_style._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans KR")
list_style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
list_style.font.size = Pt(11)
list_style.paragraph_format.left_indent = Inches(0.5)
list_style.paragraph_format.first_line_indent = Inches(-0.25)
list_style.paragraph_format.space_after = Pt(8)
list_style.paragraph_format.line_spacing = 1.167

header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
header_run = header_p.add_run("강냉이 에스크  |  실무자 피드백용 프로토타입")
set_run_font(header_run, size=9, bold=True, color=MUTED)
add_page_number(section.footer.paragraphs[0])

numbering_id = add_custom_numbering(doc)

# 1. Opening
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(14)
kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(5)
kr = kicker.add_run("PROTOTYPE PLANNING BRIEF")
set_run_font(kr, size=9.5, bold=True, color=BLUE)
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(5)
tr = title.add_run("강냉이 에스크")
set_run_font(tr, size=25, bold=True, color=NAVY)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(18)
sr = subtitle.add_run("각 부서 실무자 피드백을 위한 AI 학생지원 챗봇 프로토타입 기획서")
set_run_font(sr, size=14, color=MUTED)

add_label_value(doc, "문서 목적", "실제 담당 부서가 답변 정확성과 업무 절감 가능성을 검증할 수 있는 공통 기준 제공")
add_label_value(doc, "핵심 사용자", "강남대학교 학생 및 반복 문의를 응대하는 학내 실무자")
add_label_value(doc, "작성", "강냉이 에스크 프로젝트팀 (김소연 · 박재희 · 노성준)")
add_label_value(doc, "기준일", "2026년 7월 21일")
add_label_value(doc, "데모", "같은 교내/로컬 네트워크에서 http://192.168.0.61:8080 접속")

add_callout(
    doc,
    "한 문장 제안",
    "학생이 질문하면 공식 공지와 FAQ에서 필요한 신청 정보와 담당 연락처를 한 화면에 정리해 주어, 실무자가 반복적으로 설명하는 시간을 줄인다.",
)

add_heading(doc, "제출물 구성", 1)
add_bullet(doc, "기획 문서: 문제, MVP 범위, 사용자 흐름, 데이터와 답변 원칙")
add_bullet(doc, "실행 데모: 질문 입력, 문의 분야/FAQ 선택, 정리된 답변, 담당 부서, 원문 확인")
add_bullet(doc, "피드백 도구: 실무자별 검증 시나리오, 평가표, 수정 우선순위 기준")

# 2. Problem and outcomes
page_break(doc)
add_heading(doc, "1. 문제 정의와 목표", 1)
add_heading(doc, "1.1 해결하려는 문제", 2)
add_body(
    doc,
    "학생은 여러 공지와 학사안내에서 신청 기간, 대상, 준비 서류와 담당 부서를 직접 찾아야 한다. 실무자는 이미 공개된 내용을 전화나 방문 문의에서 반복 설명하고, 잘못 찾아온 문의를 다시 전달하는 데 시간을 사용한다.",
)
add_bullet(doc, "공지 제목만으로 필요한 절차와 준비물을 빠르게 파악하기 어렵다.")
add_bullet(doc, "같은 질문이 학기마다 반복되며 담당 부서의 단순 안내 업무가 누적된다.")
add_bullet(doc, "공지 본문에 번호가 없으면 학생이 별도 연락처 페이지를 다시 찾아야 한다.")
add_bullet(doc, "AI가 그럴듯하지만 근거 없는 답변을 만들 경우 오히려 민원 비용이 커진다.")

add_heading(doc, "1.2 목표와 검증 가설", 2)
add_callout(
    doc,
    "검증 가설",
    "공식 공지에 근거한 요약, 신청 절차, 담당 업무 연락처와 원문 링크를 함께 제공하면 학생의 자기 해결률이 높아지고 실무자의 반복 문의가 줄어든다.",
    color=GREEN,
    fill="EAF7F2",
)
add_bullet(doc, "학생은 한 화면에서 ‘누가, 언제, 무엇을, 어떻게’ 해야 하는지 이해한다.")
add_bullet(doc, "실무자는 답변 내용과 전화번호를 원문 및 공식 연락처로 검증할 수 있다.")
add_bullet(doc, "근거가 부족한 경우 추측 대신 질문 보완 또는 담당 부서 문의로 전환한다.")

add_heading(doc, "1.3 이번 피드백의 핵심 질문", 2)
add_numbered(doc, "답변에 실무자가 반드시 안내하는 정보가 빠짐없이 포함되어 있는가?", numbering_id)
add_numbered(doc, "담당 부서, 담당 업무, 담당자와 전화번호가 실제 운영 기준과 일치하는가?", numbering_id)
add_numbered(doc, "신청 절차의 순서와 표현이 학생이 바로 실행할 만큼 명확한가?", numbering_id)
add_numbered(doc, "오답 위험이 있는 질문에서 원문 확인 또는 담당자 연결로 안전하게 전환하는가?", numbering_id)

# 3. MVP
page_break(doc)
add_heading(doc, "2. MVP 범위와 우선순위", 1)
add_body(doc, "팀 논의의 P0·P1·P2 항목을 실제 제출과 피드백 목적에 맞게 다음과 같이 통합한다.")
add_table(
    doc,
    ["우선순위", "기능", "현재 데모", "실무자 확인 사항"],
    [
        ("P0", "질문 입력·문의 분야·FAQ 선택", "구현", "학생 표현으로 쉽게 시작할 수 있는가"),
        ("P0", "공식 공지 기반 핵심 답변", "구현", "대상·기간·방법·준비물·주의사항의 정확성"),
        ("P0", "담당 부서와 연락처 안내", "구현", "부서/업무/담당자/전화번호의 일치 여부"),
        ("P0", "추가 질문과 오류 안내", "구현", "정보 부족 시 추측하지 않고 안전하게 전환하는가"),
        ("P1", "원문 공지 및 다음 행동", "구현", "원문 접근성과 행동 안내의 유용성"),
        ("P1", "공지 세부사항 구조화", "부분 검증", "신청 순서·서류·예외조건의 누락 여부"),
        ("P2", "신규 공지 자동 반영", "파이프라인 준비", "운영 주기와 검수 책임 확정 필요"),
        ("P2", "문의 통계·추천·마감 알림", "제외", "실무 피드백 후 우선순위 재결정"),
    ],
    [1050, 2700, 1500, 4110],
    font_size=8.5,
)

add_heading(doc, "2.1 첫 버전에서 제외하는 기능", 2)
add_bullet(doc, "담당 부서로 즉시 전화 연결")
add_bullet(doc, "질문 원문과 개인 문의 이력의 장기 저장")
add_bullet(doc, "일정 시간 미사용 시 자동 종료")
add_bullet(doc, "개인별 맞춤 추천과 마감 알림")
add_bullet(doc, "실무자 승인 없는 FAQ 자동 게시")

add_heading(doc, "2.2 완료 기준", 2)
add_bullet(doc, "질문 입력 후 관련 답변과 담당 부서가 표시된다.")
add_bullet(doc, "빈 입력은 ‘질문을 입력하거나 문의 분야를 선택해 주세요’로 안내한다.")
add_bullet(doc, "관련 근거가 없으면 재질문 또는 담당 부서 확인을 안내한다.")
add_bullet(doc, "답변에서 공식 원문으로 이동할 수 있다.")
add_bullet(doc, "프로덕션 빌드와 자동 테스트가 오류 없이 통과한다.")

# 4. UX
add_heading(doc, "3. 사용자 경험과 화면 구성", 1)
add_heading(doc, "3.1 기본 사용자 흐름", 2)
numbering_id = add_custom_numbering(doc)
for text in [
    "학교 홈페이지 또는 공유된 데모 주소에서 ‘강냉이 에스크’를 실행한다.",
    "질문을 직접 입력하거나 문의 분야·자주 묻는 질문을 선택한다.",
    "서비스가 공지와 FAQ를 검색해 핵심 답변을 구성한다.",
    "학생은 기간, 신청 방법, 준비물, 담당 부서와 관련 공지를 확인한다.",
    "정보가 충분하면 절차를 진행하고, 부족하면 추가 질문 또는 원문·담당 부서 확인으로 이동한다.",
]:
    add_numbered(doc, text, numbering_id)

add_heading(doc, "3.2 화면 상태별 요구사항", 2)
add_table(
    doc,
    ["상태", "화면 요소", "예상 행동"],
    [
        ("입력 전", "안내 문구, 질문창, 분야 버튼, FAQ/예시 질문", "질문 또는 항목 선택"),
        ("답변", "요약, 대상, 기간, 방법, 서류, 담당 부서, 원문", "절차 진행 또는 원문 확인"),
        ("정보 부족", "부족한 근거와 질문 보완 안내", "학년도·학기·대상 추가"),
        ("오류", "관련 정보를 찾지 못했다는 안내", "재질문 또는 담당 부서 확인"),
    ],
    [1200, 4770, 3390],
)

add_heading(doc, "3.3 답변 화면 예시", 2)
if SCREENSHOT.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(SCREENSHOT), width=Inches(6.35))
    doc_pr = run._r.xpath(".//wp:docPr")[0]
    doc_pr.set("title", "강냉이 에스크 데모 답변 화면")
    doc_pr.set("descr", "2026학년도 2학기 수강신청 일정을 질문한 뒤 단계별 답변과 담당 부서가 표시된 화면")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)
    cr = caption.add_run("그림 1. 수강신청 일정 질문에 대한 공식 공지 기반 답변")
    set_run_font(cr, size=9, italic=True, color=MUTED)

# 5. Response/data
add_heading(doc, "4. 답변 원칙과 데이터 구성", 1)
add_heading(doc, "4.1 답변에 반드시 포함할 정보", 2)
for text in [
    "핵심 요약: 학생이 지금 알아야 할 결론",
    "신청 대상과 제외 조건",
    "본 신청 기간과 예비·변경 등 중요 일정",
    "준비 서류, 비용, 장소와 신청 방식",
    "실행 가능한 처리 순서(1단계부터 마지막 확인까지)",
    "담당 부서, 담당 업무, 담당자, 전화번호",
    "공식 원문 링크와 확인 기준 시각",
]:
    add_bullet(doc, text)

add_heading(doc, "4.2 근거 우선순위", 2)
numbering_id = add_custom_numbering(doc)
add_numbered(doc, "공지 원문과 첨부파일에 명시된 값", numbering_id)
add_numbered(doc, "학사안내와 공식 FAQ의 상시 기준", numbering_id)
add_numbered(doc, "학교 공식 직원 연락처의 부서·업무별 연락처", numbering_id)
add_numbered(doc, "근거가 확인되지 않은 값은 생성하지 않고 ‘확인 필요’로 전환", numbering_id)

add_heading(doc, "4.3 현재 프로토타입 데이터 흐름", 2)
add_body(
    doc,
    "학교 홈페이지의 공지·학사안내·FAQ·직원 연락처를 수집해 원문과 추출 결과를 저장한다. 긴 문서는 검색 가능한 단위로 나누고 의미 기반 검색 인덱스를 만든다. 질문 시에는 관련 문서를 찾은 뒤 확인된 구조화 정보만 답변 카드에 표시하며, 원문 링크를 함께 제공한다.",
)
add_table(
    doc,
    ["단계", "처리", "안전장치"],
    [
        ("수집", "HTML·첨부·이미지 OCR 저장", "원문과 추출 상태 보존"),
        ("정리", "기간·대상·절차·연락처 구조화", "원문에 없는 값 제거"),
        ("검색", "키워드+의미 검색으로 후보 선정", "분류·학기·현재성 반영"),
        ("답변", "확정 값은 카드, 설명은 짧은 문장", "근거 부족 시 추측 금지"),
        ("개선", "아쉬운 답변과 누락 필드 집계", "원문 질문·개인정보 저장 최소화"),
    ],
    [1050, 4380, 3930],
)

add_callout(
    doc,
    "연락처 보완 원칙",
    "공지에 전화번호가 없을 때만 공식 직원 연락처에서 ‘부서 + 담당 업무’를 함께 대조한다. 예: 교무팀 + 수강신청 → 해당 업무 담당자 번호. 화면에는 공식 연락처에서 보완한 값임을 표시한다.",
    color=GOLD,
    fill="FFF8E8",
)

# 6. Demo guide
page_break(doc)
add_heading(doc, "5. 데모 실행 가이드", 1)
add_heading(doc, "5.1 접속", 2)
p = doc.add_paragraph()
label = p.add_run("내부 데모 주소  ")
set_run_font(label, bold=True, color=NAVY)
add_hyperlink(p, "http://192.168.0.61:8080", "http://192.168.0.61:8080")
add_body(doc, "현재 주소는 데모 서버와 같은 네트워크에서 접속할 수 있다. 외부 실무자 배포 시에는 고정 도메인의 Cloudflare Named Tunnel로 전환하고 동일한 문서에 최종 주소를 갱신한다.")

add_heading(doc, "5.2 권장 시연 순서(약 7분)", 2)
numbering_id = add_custom_numbering(doc)
for text in [
    "첫 화면에서 문의 분야·FAQ·질문 입력의 세 가지 진입 방식을 보여준다.",
    "‘2026학년도 2학기 수강신청 일정 알려줘’를 입력해 일정 카드와 담당자·원문을 확인한다.",
    "‘휴학 신청 방법 알려줘’를 입력해 순서형 처리 절차와 담당 업무 연락처를 확인한다.",
    "존재하지 않는 질문을 입력해 근거 부족 시 추측하지 않는 동작을 확인한다.",
    "‘아쉬워요’를 선택해 실무자가 수정 필요 사유를 남기는 흐름을 설명한다.",
]:
    add_numbered(doc, text, numbering_id)

add_heading(doc, "5.3 실무자 검증용 질문 세트", 2)
add_table(
    doc,
    ["분야", "대표 질문", "확인할 핵심"],
    [
        ("교무", "2학기 수강신청 일정과 방법 알려줘", "본 신청/예비 일정, 시스템, 담당자"),
        ("학적", "일반휴학 신청 순서와 준비 서류 알려줘", "메뉴 경로, 승인 절차, 예외조건"),
        ("장학", "국가장학금 신청 대상과 기간 알려줘", "학기·차수, 대상, 외부 신청 링크"),
        ("등록", "등록금 납부 기간과 분할납부 방법 알려줘", "기간, 계좌/시스템, 담당 부서"),
        ("학생지원", "예비군 신고는 어디에 제출해?", "대상, 제출처, 담당 연락처"),
        ("오류 대응", "교내 수영장 수온 알려줘", "근거 없음 안내와 안전한 전환"),
    ],
    [1200, 4530, 3630],
    font_size=9,
)

# 7. Feedback protocol
page_break(doc)
add_heading(doc, "6. 실무자 피드백 프로토콜", 1)
add_heading(doc, "6.1 참여 부서와 역할", 2)
add_table(
    doc,
    ["참여자", "권장 부서", "주요 검증 책임"],
    [
        ("업무 담당자", "교무팀·장학복지팀·재무회계팀 등", "내용, 절차, 일정, 연락처 정확성"),
        ("상담/민원 담당", "학생지원·대학일자리 등", "반복 문의 감소 가능성과 표현 이해도"),
        ("학생 사용자", "재학생 3~5명", "질문 입력과 다음 행동의 사용성"),
        ("프로젝트팀", "기획·개발", "오류 재현, 데이터 누락 기록, 개선 반영"),
    ],
    [1500, 3060, 4800],
)

add_heading(doc, "6.2 1인당 진행 방식(15~20분)", 2)
numbering_id = add_custom_numbering(doc)
for text in [
    "사전 설명 2분: 서비스 목적과 ‘공지 대체가 아닌 안내 보조’ 원칙을 설명한다.",
    "자유 질문 5분: 담당자가 실제로 자주 받는 질문 3개를 직접 입력한다.",
    "정확성 검토 5분: 답변, 담당자, 전화번호와 원문을 대조한다.",
    "평가표 작성 3분: 항목별 1~5점과 수정 문구를 기록한다.",
    "인터뷰 5분: 줄어들 문의, 위험한 답변, 추가로 필요한 데이터를 묻는다.",
]:
    add_numbered(doc, text, numbering_id)

add_heading(doc, "6.3 반드시 기록할 피드백", 2)
add_bullet(doc, "질문 원문 대신 질문 유형과 담당 업무(개인정보 제외)")
add_bullet(doc, "잘못되거나 오래된 사실과 정확한 수정값")
add_bullet(doc, "빠진 필드: 담당자, 전화번호, 기간, 대상, 서류, 링크, 예외조건 등")
add_bullet(doc, "원문은 맞지만 학생에게 오해를 줄 수 있는 표현")
add_bullet(doc, "챗봇이 해결해도 되는 문의와 반드시 사람에게 넘겨야 하는 문의")

add_heading(doc, "6.4 평가 지표와 합격 기준", 2)
add_table(
    doc,
    ["지표", "측정 방법", "프로토타입 합격 기준"],
    [
        ("근거 적합성", "관련 공지/FAQ가 실제 질문과 일치", "90% 이상"),
        ("핵심 정보 정확성", "기간·대상·서류·방법 대조", "중대 오류 0건"),
        ("연락처 정확성", "부서·업무·담당자·번호 대조", "100% 또는 미제공"),
        ("절차 실행 가능성", "학생이 추가 설명 없이 다음 단계 선택", "평균 4/5 이상"),
        ("원문 추적성", "모든 주요 답변에서 원문 접근", "100%"),
        ("안전한 실패", "근거 부족 질문에서 추측 여부", "추측 답변 0건"),
    ],
    [2100, 4260, 3000],
    font_size=9,
)

# 8. Feedback form
add_heading(doc, "7. 실무자 피드백 기록지", 1)
add_body(doc, "아래 기록지는 한 질문당 한 장을 기준으로 사용한다. 개인정보와 학생 개별 사례는 기록하지 않는다.")
add_label_value(doc, "부서 / 담당 업무", "________________________________________________________")
add_label_value(doc, "검토한 질문", "________________________________________________________")
add_label_value(doc, "기준 원문 또는 내부 기준", "________________________________________________________")

add_heading(doc, "7.1 항목별 평가", 2)
add_table(
    doc,
    ["평가 항목", "1", "2", "3", "4", "5", "비고"],
    [
        ("질문과 답변의 관련성", "□", "□", "□", "□", "□", ""),
        ("기간·대상·방법 정확성", "□", "□", "□", "□", "□", ""),
        ("담당 부서·연락처 정확성", "□", "□", "□", "□", "□", ""),
        ("처리 순서의 이해 용이성", "□", "□", "□", "□", "□", ""),
        ("원문 확인의 편의성", "□", "□", "□", "□", "□", ""),
    ],
    [2550, 540, 540, 540, 540, 540, 4110],
    font_size=9,
)

add_heading(doc, "7.2 오류·누락 기록", 2)
for label in [
    "잘못된 내용",
    "빠진 정보",
    "권장 수정 문구",
    "사람에게 연결해야 하는 조건",
]:
    add_label_value(doc, label, "____________________________________________________________")

add_heading(doc, "7.3 최종 판단", 2)
add_bullet(doc, "□ 현재 상태로 사용 가능")
add_bullet(doc, "□ 경미한 수정 후 사용 가능")
add_bullet(doc, "□ 데이터 또는 답변 구조 보완 후 재검토 필요")
add_bullet(doc, "□ 해당 문의는 챗봇 안내 대상에서 제외")

# 9. Current status and next steps
page_break(doc)
add_heading(doc, "8. 현재 구현 상태와 다음 단계", 1)
add_heading(doc, "8.1 현재 데모 기준", 2)
add_bullet(doc, "공식 공지 757건과 공식 직원 연락처 373건을 검색·업무별 연락처 데이터로 사용한다.")
add_bullet(doc, "대표 공지에서 이미지 OCR, 일정·절차·담당자 추출과 기간·방법·연락처 카드 표시를 검증했다.")
add_bullet(doc, "공식 원문 연결을 제공하며 백엔드 테스트 62개와 프런트엔드 테스트 7개가 통과했다.")

add_heading(doc, "8.2 현재 제한사항", 2)
add_bullet(doc, "전체 공지를 최종 품질로 재구조화한 상태는 아니며 대표 범위에서 구조를 검증 중이다.")
add_bullet(doc, "자동 수집 스케줄은 운영 승인 전까지 비활성화되어 있으며 갱신·검수 책임자를 정해야 한다.")
add_bullet(doc, "현재 주소는 동일 네트워크용이므로 외부 피드백 전 고정 터널 주소가 필요하다.")

add_heading(doc, "8.3 피드백 반영 우선순위", 2)
add_table(
    doc,
    ["등급", "판정 기준", "처리"],
    [
        ("즉시 수정", "잘못된 기간·대상·전화번호, 근거 없는 답변", "공개 차단 후 데이터 수정"),
        ("다음 배포", "절차 누락, 불명확한 표현, 잘못된 추천 공지", "프롬프트·검색·구조화 보완"),
        ("후속 검토", "맞춤 추천, 알림, 통계 등 확장 요구", "효과·비용 평가 후 로드맵 반영"),
    ],
    [1500, 4800, 3060],
)

add_heading(doc, "8.4 제출 전 체크리스트", 2)
numbering_id = add_custom_numbering(doc)
for text in [
    "외부 접속용 고정 데모 주소를 확인하고 교무·장학·등록 등 최소 3개 부서의 대표 질문을 사전 점검한다.",
    "잘못된 연락처 또는 기간이 있는 답변은 데모에서 제외한다.",
    "피드백 기록지와 진행 방식을 전달하고, 각 부서는 대표 질문 3개와 기준 답변·연락처를 제공한다. 결과는 즉시 수정/다음 배포/후속 검토로 분류한다.",
]:
    add_numbered(doc, text, numbering_id)

# Keep tables visually grouped with the preceding paragraph where possible.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if not run.font.name:
            set_run_font(run)

doc.core_properties.title = "강냉이 에스크 실무자 피드백용 프로토타입 기획서"
doc.core_properties.subject = "AI 학생지원 챗봇 프로토타입 및 실무자 피드백 계획"
doc.core_properties.author = "강냉이 에스크 프로젝트팀"
doc.core_properties.keywords = "강냉이 에스크, 강남대학교, AI 챗봇, 프로토타입, 실무자 피드백"
doc.save(OUTPUT)
print(OUTPUT)
